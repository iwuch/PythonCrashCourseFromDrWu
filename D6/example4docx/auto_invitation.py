from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
name = 'kobe bryant'
doc = Document()
doc.add_heading('invitation',level =1)
paragraph1 = doc.add_paragraph(name.title())
paragraph2 = doc.add_paragraph('Welcome to my party.')
paragraph3 = doc.add_paragraph()
paragraph3.add_run('Thank you!').bold = True
paragraph3.add_run('Best regards from')
paragraph3.add_run('Dr.Wu').italic = True
doc.add_page_break()
doc.save('invitation_'+name+'.docx')